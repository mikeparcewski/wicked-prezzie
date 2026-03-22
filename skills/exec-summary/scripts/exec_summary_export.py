#!/usr/bin/env python3
"""Export a markdown executive summary to a formatted Word (.docx) document.

Reads an exec-summary.md file (8-section schema), parses each section by
heading, and produces a professionally formatted Word document with title page,
table of contents, heading styles, page numbers, and optional draft watermark.

Usage:
    python exec_summary_export.py exec-summary.md --output exec-summary.docx
    python exec_summary_export.py exec-summary.md --output draft.docx --draft
    python exec_summary_export.py exec-summary.md --output out.docx --title "Project Alpha"
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.text.paragraph import Paragraph

_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(_root / "shared"))
from paths import output_path, ensure_output_dir


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)


def parse_markdown_sections(md_text: str) -> list[dict]:
    """Parse markdown into a list of sections with heading level, title, and body.

    Returns:
        List of dicts: {"level": int, "title": str, "body": str}
    """
    matches = list(_HEADING_RE.finditer(md_text))
    if not matches:
        return [{"level": 0, "title": "", "body": md_text.strip()}]

    sections = []
    # Capture any content before the first heading as preamble.
    preamble = md_text[: matches[0].start()].strip()
    if preamble:
        sections.append({"level": 0, "title": "", "body": preamble})

    for i, m in enumerate(matches):
        level = len(m.group(1))
        title = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md_text)
        body = md_text[start:end].strip()
        sections.append({"level": level, "title": title, "body": body})

    return sections


# ---------------------------------------------------------------------------
# Inline markdown → runs
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*"   # bold-italic
    r"|\*\*(.+?)\*\*"        # bold
    r"|\*(.+?)\*"            # italic
    r"|`(.+?)`"              # inline code
    r"|([^*`]+))"            # plain text
)


def _add_runs(paragraph: Paragraph, text: str, font_name: str, font_size: Pt):
    """Add runs to a paragraph with basic inline markdown formatting."""
    for m in _INLINE_RE.finditer(text):
        bold_italic = m.group(2)
        bold = m.group(3)
        italic = m.group(4)
        code = m.group(5)
        plain = m.group(6)

        if bold_italic:
            run = paragraph.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold:
            run = paragraph.add_run(bold)
            run.bold = True
        elif italic:
            run = paragraph.add_run(italic)
            run.italic = True
        elif code:
            run = paragraph.add_run(code)
            run.font.name = "Consolas"
            run.font.size = Pt(font_size.pt - 1) if hasattr(font_size, "pt") else Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            continue
        else:
            run = paragraph.add_run(plain)

        run.font.name = font_name
        run.font.size = font_size


# ---------------------------------------------------------------------------
# Document styling
# ---------------------------------------------------------------------------

FONT_NAME = "Calibri"
BODY_SIZE = Pt(11)
H1_SIZE = Pt(22)
H2_SIZE = Pt(16)
H3_SIZE = Pt(13)
TITLE_SIZE = Pt(32)
SUBTITLE_SIZE = Pt(14)

HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)  # Dark navy
ACCENT_COLOR = RGBColor(0x2E, 0x5C, 0x8A)   # Steel blue


def _style_heading(paragraph: Paragraph, level: int):
    """Apply consistent heading styling."""
    size = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}.get(level, H3_SIZE)
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.color.rgb = HEADING_COLOR
        run.bold = True
    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(8)


def _add_page_numbers(section):
    """Add page numbers centered in the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # PAGE field
    run = p.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def _add_watermark(section, text: str = "DRAFT \u2014 For Review"):
    """Add a diagonal text watermark to the section header.

    Uses raw lxml elements for VML shapes since python-docx does not register
    the VML namespace prefix.
    """
    from lxml import etree

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Build VML watermark shape as raw XML — the only reliable way to create
    # VML elements with python-docx, which does not map the v: namespace.
    watermark_xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '        xmlns:v="urn:schemas-microsoft-com:vml"'
        '        xmlns:o="urn:schemas-microsoft-com:office:office">'
        '  <v:shape id="WatermarkShape" type="#_x0000_t136"'
        '    style="position:absolute;margin-left:0;margin-top:0;'
        '           width:500pt;height:100pt;rotation:315;'
        '           z-index:-251658752;'
        '           mso-position-horizontal:center;'
        '           mso-position-horizontal-relative:margin;'
        '           mso-position-vertical:center;'
        '           mso-position-vertical-relative:margin"'
        '    fillcolor="#C0C0C0" stroked="f">'
        '    <v:fill opacity=".3"/>'
        f'    <v:textpath style=\'font-family:"Calibri";font-size:1pt\''
        f'      string="{text}"/>'
        '  </v:shape>'
        '</w:pict>'
    )
    pict_elem = etree.fromstring(watermark_xml)
    run = p.add_run()
    run._element.append(pict_elem)


def _add_toc(doc: Document):
    """Insert a table of contents field that updates on open in Word."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)

    # TOC title
    title_run = p.add_run("Table of Contents")
    title_run.bold = True
    title_run.font.name = FONT_NAME
    title_run.font.size = H2_SIZE
    title_run.font.color.rgb = HEADING_COLOR

    # Instruction paragraph with TOC field
    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    run2 = toc_p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instr)

    run3 = toc_p.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_separate)

    run4 = toc_p.add_run("[Right-click and select 'Update Field' to populate]")
    run4.font.name = FONT_NAME
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = toc_p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_end)

    # Page break after TOC
    doc.add_page_break()


def _add_title_page(doc: Document, title: str):
    """Add a title page with the deck title centered vertically."""
    # Add several blank paragraphs to push title toward center
    for _ in range(8):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.font.name = FONT_NAME
    run.font.size = TITLE_SIZE
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    title_p.paragraph_format.space_after = Pt(16)

    # Accent line
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent_run = line_p.add_run("\u2500" * 40)
    accent_run.font.color.rgb = ACCENT_COLOR
    accent_run.font.size = Pt(14)
    line_p.paragraph_format.space_after = Pt(16)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Executive Summary")
    sub_run.font.name = FONT_NAME
    sub_run.font.size = SUBTITLE_SIZE
    sub_run.font.color.rgb = ACCENT_COLOR
    sub_run.italic = True

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Body content rendering
# ---------------------------------------------------------------------------

_BULLET_RE = re.compile(r"^[\s]*[-*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^[\s]*\d+[.)]\s+(.+)$")
_HR_RE = re.compile(r"^---+\s*$")


def _render_body(doc: Document, body: str):
    """Render the body text of a section into the document.

    Handles paragraphs, bullet lists, numbered lists, and horizontal rules.
    """
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule — skip (section separators are handled by headings)
        if _HR_RE.match(line):
            i += 1
            continue

        # Sub-heading within body (## or ### that wasn't split into its own section)
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(
                level, "Heading 3"
            )
            p = doc.add_paragraph(heading_match.group(2).strip(), style=style)
            _style_heading(p, level)
            i += 1
            continue

        # Bullet list
        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, bullet_match.group(1), FONT_NAME, BODY_SIZE)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered list
        num_match = _NUMBERED_RE.match(line)
        if num_match:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, num_match.group(1), FONT_NAME, BODY_SIZE)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Regular paragraph — collect consecutive non-blank, non-special lines
        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if not next_line.strip():
                break
            if _HR_RE.match(next_line):
                break
            if re.match(r"^#{1,6}\s+", next_line):
                break
            if _BULLET_RE.match(next_line) or _NUMBERED_RE.match(next_line):
                break
            para_lines.append(next_line.strip())
            i += 1

        text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_runs(p, text, FONT_NAME, BODY_SIZE)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(16)


# ---------------------------------------------------------------------------
# Main export
# ---------------------------------------------------------------------------


def export_exec_summary(
    md_path: str,
    output_path_str: str,
    title: str = "Executive Summary",
    draft: bool = False,
) -> str:
    """Export a markdown exec-summary to a formatted Word document.

    Args:
        md_path: Path to the source exec-summary.md file.
        output_path_str: Path for the output .docx file.
        title: Title for the title page.
        draft: If True, add a "DRAFT - For Review" watermark.

    Returns:
        Absolute path to the created .docx file.
    """
    md_file = Path(md_path)
    if not md_file.exists():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    md_text = md_file.read_text(encoding="utf-8")
    sections = parse_markdown_sections(md_text)

    doc = Document()

    # -- Document defaults --
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Configure heading styles
    for level, size in [(1, H1_SIZE), (2, H2_SIZE), (3, H3_SIZE)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = FONT_NAME
        h_style.font.size = size
        h_style.font.color.rgb = HEADING_COLOR
        h_style.font.bold = True
        h_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        h_style.paragraph_format.space_after = Pt(8)

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Page numbers
    _add_page_numbers(section)

    # Watermark
    if draft:
        _add_watermark(section)

    # -- Title page --
    _add_title_page(doc, title)

    # -- Table of contents --
    _add_toc(doc)

    # -- Sections --
    for sec in sections:
        level = sec["level"]
        sec_title = sec["title"]
        body = sec["body"]

        if level == 0:
            # Preamble text (before any heading)
            if body:
                _render_body(doc, body)
            continue

        # Add heading
        heading_style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(
            level, "Heading 3"
        )
        h_para = doc.add_paragraph(sec_title, style=heading_style)
        _style_heading(h_para, level)

        # Add body content
        if body:
            _render_body(doc, body)

    # -- Save --
    out = Path(output_path_str)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return str(out.resolve())


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Export an exec-summary markdown file to Word (.docx).",
    )
    parser.add_argument(
        "input",
        help="Path to the exec-summary.md file.",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Output .docx path. Defaults to <input-stem>.docx in the output directory.",
    )
    parser.add_argument(
        "--title", "-t",
        default="Executive Summary",
        help="Title for the cover page (default: 'Executive Summary').",
    )
    parser.add_argument(
        "--draft",
        action="store_true",
        help='Add a "DRAFT - For Review" watermark to every page.',
    )

    args = parser.parse_args()

    if args.output is None:
        ensure_output_dir()
        stem = Path(args.input).stem
        args.output = output_path(f"{stem}.docx")

    result = export_exec_summary(
        md_path=args.input,
        output_path_str=args.output,
        title=args.title,
        draft=args.draft,
    )
    print(f"Exported: {result}")


if __name__ == "__main__":
    main()
