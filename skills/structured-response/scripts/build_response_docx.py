#!/usr/bin/env python3
"""
build_response_docx.py — Assemble section markdown files into a Word document.

Reads response-config.md for structure, sections/{id}/final.md for content,
and sections/{id}/comments.json for inline reviewer comments.

Produces a .docx with proper heading hierarchy, styled body text, inline
comment bubbles, and highlighted PLACEHOLDER markers.

Dependencies: python-docx, pyyaml
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is required: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("pyyaml is required: pip install pyyaml", file=sys.stderr)
    sys.exit(1)


def parse_config(config_path: Path) -> dict:
    """Parse response-config.md — extract YAML frontmatter and section defs."""
    text = config_path.read_text(encoding="utf-8")

    # Extract YAML frontmatter
    fm_match = re.match(r"^---\n(.*?)\n---", text, re.DOTALL)
    frontmatter = yaml.safe_load(fm_match.group(1)) if fm_match else {}

    # Extract sections by parsing markdown headers
    sections = []
    section_pattern = re.compile(
        r"^###\s+\d+\.\s+(.+?)$\n(.*?)(?=^###\s+\d+\.|^##\s+Reviewers|$)",
        re.MULTILINE | re.DOTALL,
    )
    for match in section_pattern.finditer(text):
        title = match.group(1).strip()
        body = match.group(2).strip()
        section_id = None
        page_limit = None
        for line in body.splitlines():
            line_s = line.strip()
            if line_s.startswith("- **id**:"):
                section_id = line_s.split(":", 1)[1].strip()
            elif line_s.startswith("- **page_limit**:"):
                val = line_s.split(":", 1)[1].strip()
                page_limit = int(val) if val.isdigit() else None
        if section_id:
            sections.append({
                "id": section_id,
                "title": title,
                "page_limit": page_limit,
            })

    frontmatter["sections"] = sections
    return frontmatter


def add_comment(paragraph, text: str, author: str, initials: str, doc):
    """Add an inline comment to a paragraph using OOXML."""
    # Get or create comments part
    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        # Create comments element directly on the paragraph
        # Simplified: add comment as a footnote-style reference
        run = paragraph.add_run(f" [{author}: {text}]")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        run.font.italic = True
        return

    comment_el = OxmlElement("w:comment")
    comment_el.set(qn("w:id"), str(id(text) % 10000))
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:initials"), initials)
    comment_el.set(qn("w:date"), datetime.now().isoformat())

    comment_p = OxmlElement("w:p")
    comment_r = OxmlElement("w:r")
    comment_t = OxmlElement("w:t")
    comment_t.text = text
    comment_r.append(comment_t)
    comment_p.append(comment_r)
    comment_el.append(comment_p)


def highlight_placeholders(paragraph):
    """Find [PLACEHOLDER: ...] in paragraph text and highlight yellow."""
    for run in paragraph.runs:
        if "[PLACEHOLDER:" in run.text:
            run.font.highlight_color = 7  # Yellow


def markdown_to_docx(doc, md_text: str, base_heading_level: int = 2):
    """Convert simple markdown to docx paragraphs.

    Handles: headings (##, ###, ####), bullet lists (- ), numbered lists,
    bold (**text**), italic (*text*), and body paragraphs.
    """
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            # Map markdown heading level to Word heading level
            word_level = min(level + base_heading_level - 2, 9)
            doc.add_heading(text, level=word_level)
            i += 1
            continue

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text)
            highlight_placeholders(p)
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text)
            highlight_placeholders(p)
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not lines[i].strip().startswith("- ") and not lines[i].strip().startswith("* ") and not re.match(r"^\d+\.\s+", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1

        text = " ".join(para_lines)
        p = doc.add_paragraph()
        _add_inline_formatting(p, text)
        highlight_placeholders(p)


def _add_inline_formatting(paragraph, text: str):
    """Parse bold (**text**) and italic (*text*) inline markers."""
    # Split on bold/italic markers
    pattern = r"(\*\*.*?\*\*|\*.*?\*)"
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def build_document(config_path: Path, sections_dir: Path, output_path: Path):
    """Main assembly: read config, read sections, build Word document."""
    config = parse_config(config_path)
    doc = Document()

    # Title
    title = config.get("document_name", "Document")
    doc.add_heading(title, level=0)

    # Process each section
    for section in config.get("sections", []):
        section_id = section["id"]
        section_title = section["title"]
        final_path = sections_dir / section_id / "final.md"
        comments_path = sections_dir / section_id / "comments.json"

        # Section heading
        doc.add_heading(section_title, level=1)

        if final_path.exists():
            md_text = final_path.read_text(encoding="utf-8")
            markdown_to_docx(doc, md_text)
        else:
            p = doc.add_paragraph(f"[Section not yet generated: {section_id}]")
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Add comments as styled annotations
        if comments_path.exists():
            comments = json.loads(comments_path.read_text(encoding="utf-8"))
            if comments:
                # Add a subtle separator
                doc.add_paragraph()
                for comment in comments:
                    reviewer = comment.get("reviewer", "Reviewer")
                    finding = comment.get("finding", "")
                    comment_type = comment.get("type", "comment")
                    initials = "".join(
                        w[0].upper() for w in reviewer.split() if w
                    )[:3]

                    # For now, add as styled inline annotations
                    # Full OOXML comment support requires document template setup
                    p = doc.add_paragraph()
                    prefix = p.add_run(f"[{reviewer}] ")
                    prefix.font.size = Pt(8)
                    prefix.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                    prefix.bold = True

                    body = p.add_run(finding)
                    body.font.size = Pt(8)
                    body.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                    body.italic = True

        # Page break between sections (except last)
        if section != config["sections"][-1]:
            doc.add_page_break()

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Document saved to: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Assemble section markdown into a Word document"
    )
    parser.add_argument(
        "--config",
        required=True,
        help="Path to response-config.md",
    )
    parser.add_argument(
        "--sections-dir",
        required=True,
        help="Path to sections/ directory",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output .docx path",
    )
    args = parser.parse_args()

    config_path = Path(args.config)
    sections_dir = Path(args.sections_dir)
    output_path = Path(args.output)

    if not config_path.exists():
        print(f"Config not found: {config_path}", file=sys.stderr)
        sys.exit(1)
    if not sections_dir.exists():
        print(f"Sections directory not found: {sections_dir}", file=sys.stderr)
        sys.exit(1)

    build_document(config_path, sections_dir, output_path)


if __name__ == "__main__":
    main()
