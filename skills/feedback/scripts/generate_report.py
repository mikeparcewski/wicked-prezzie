"""Generate a formatted feedback report from analysis results.

Produces markdown or structured output suitable for Claude to present
or for further processing. Can also generate a Word (.docx) report.

Usage:
    python generate_report.py report.json [--format markdown|docx] [--output report.md]

Or programmatically:
    from generate_report import render_markdown, render_docx
"""

import json
import sys
from pathlib import Path


def render_markdown(report):
    """Render analysis report as markdown."""
    lines = []
    summary = report["summary"]

    lines.append("# Feedback Analysis Report")
    lines.append("")
    lines.append("## Overview")
    lines.append("")
    lines.append(f"- **{summary['total_comments']}** comments from **{summary['reviewer_count']}** reviewers")
    lines.append(f"- Reviewers: {', '.join(summary['reviewers'])}")
    lines.append(f"- Sections with feedback: {summary['sections_with_feedback']}/{summary['total_sections']}")
    lines.append(f"- Endorsements: {summary['endorsements']} | Concerns: {summary['concerns']} | Suggestions: {summary['suggestions']}")
    lines.append(f"- **{summary['alignment_points']}** alignment points | **{summary['divergence_points']}** divergence points")
    lines.append(f"- **{summary['high_priority_actions']}** high-priority actions")
    lines.append("")

    # Narrative implications
    if report.get("narrative_implications"):
        lines.append("## What This Means for the Narrative")
        lines.append("")
        for impl in report["narrative_implications"]:
            icon = {
                "attention_hotspot": "🔥",
                "silent_sections": "🤫",
                "consensus_for_change": "✅",
                "requires_discussion": "⚠️",
                "overall_direction": "🧭",
            }.get(impl["type"], "📌")
            lines.append(f"**{icon} {impl['type'].replace('_', ' ').title()}**")
            lines.append(f"> {impl['detail']}")
            lines.append("")

    # Action items
    if report.get("action_items"):
        lines.append("## Action Items (Prioritized)")
        lines.append("")
        for i, item in enumerate(report["action_items"], 1):
            priority_tag = "🔴 HIGH" if item["priority"] == "high" else "🟡 MEDIUM"
            lines.append(f"### {i}. [{priority_tag}] {item['action']}")
            lines.append(f"**Section:** {item['section']}")
            if item.get("referenced_text"):
                lines.append(f"**Referenced text:** \"{item['referenced_text']}\"")
            lines.append(f"**Feedback:** {item['feedback_summary']}")
            lines.append(f"**Reviewers involved:** {item['reviewer_count']}")
            lines.append("")

    # Alignment detail
    if report.get("alignment"):
        lines.append("## Where Reviewers Agree")
        lines.append("")
        for point in report["alignment"]:
            lines.append(f"### {point['section']}")
            lines.append(f"**Type:** {point['type'].replace('_', ' ')}")
            lines.append(f"**Reviewers:** {', '.join(point['reviewers'])}")
            if point.get("referenced_text"):
                lines.append(f"**On:** \"{point['referenced_text'][:150]}\"")
            for comment in point["comments"]:
                lines.append(f"- {comment}")
            lines.append("")

    # Divergence detail
    if report.get("divergence"):
        lines.append("## Where Reviewers Diverge")
        lines.append("")
        for point in report["divergence"]:
            lines.append(f"### {point['section']}")
            if point.get("referenced_text"):
                lines.append(f"**On:** \"{point['referenced_text'][:150]}\"")
            lines.append("")
            lines.append("| Reviewer | Position | Comment |")
            lines.append("|----------|----------|---------|")
            for pos in point["positions"]:
                lines.append(f"| {pos['author']} | {pos['sentiment']} | {pos['text'][:80]} |")
            lines.append("")
            lines.append("*This requires a discussion before revisions can proceed.*")
            lines.append("")

    # Per-section breakdown
    if report.get("section_analysis"):
        lines.append("## Section-by-Section Breakdown")
        lines.append("")
        for section in report["section_analysis"]:
            if section["total_comments"] == 0:
                continue
            sb = section["sentiment_breakdown"]
            lines.append(f"### {section['section']}")
            lines.append(f"- Comments: {section['total_comments']} from {', '.join(section['reviewers'])}")
            lines.append(f"- Sentiment: ✅{sb['endorsements']} ⚠️{sb['concerns']} 💡{sb['suggestions']} 👁️{sb['observations']}")
            lines.append("")

    return "\n".join(lines)


def render_docx(report, output_path):
    """Render analysis report as a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()
    summary = report["summary"]

    # Title
    title = doc.add_heading("Feedback Analysis Report", level=0)

    # Overview
    doc.add_heading("Overview", level=1)
    overview_items = [
        f"{summary['total_comments']} comments from {summary['reviewer_count']} reviewers",
        f"Reviewers: {', '.join(summary['reviewers'])}",
        f"Sections with feedback: {summary['sections_with_feedback']}/{summary['total_sections']}",
        f"Endorsements: {summary['endorsements']} | Concerns: {summary['concerns']} | Suggestions: {summary['suggestions']}",
        f"{summary['alignment_points']} alignment points | {summary['divergence_points']} divergence points",
        f"{summary['high_priority_actions']} high-priority actions",
    ]
    for item in overview_items:
        doc.add_paragraph(item, style="List Bullet")

    # Narrative implications
    if report.get("narrative_implications"):
        doc.add_heading("What This Means for the Narrative", level=1)
        for impl in report["narrative_implications"]:
            p = doc.add_paragraph()
            run = p.add_run(impl["type"].replace("_", " ").title() + ": ")
            run.bold = True
            p.add_run(impl["detail"])

    # Action items
    if report.get("action_items"):
        doc.add_heading("Action Items (Prioritized)", level=1)
        for i, item in enumerate(report["action_items"], 1):
            priority = "HIGH" if item["priority"] == "high" else "MEDIUM"
            doc.add_heading(f"{i}. [{priority}] {item['action']}", level=2)

            p = doc.add_paragraph()
            run = p.add_run("Section: ")
            run.bold = True
            p.add_run(item["section"])

            if item.get("referenced_text"):
                p = doc.add_paragraph()
                run = p.add_run("Referenced text: ")
                run.bold = True
                p.add_run(f'"{item["referenced_text"]}"')

            p = doc.add_paragraph()
            run = p.add_run("Feedback: ")
            run.bold = True
            p.add_run(item["feedback_summary"])

    # Divergence
    if report.get("divergence"):
        doc.add_heading("Where Reviewers Diverge", level=1)
        for point in report["divergence"]:
            doc.add_heading(point["section"], level=2)
            if point.get("referenced_text"):
                p = doc.add_paragraph()
                run = p.add_run("On: ")
                run.bold = True
                p.add_run(f'"{point["referenced_text"][:150]}"')

            # Table for positions
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Reviewer"
            hdr[1].text = "Position"
            hdr[2].text = "Comment"
            for pos in point["positions"]:
                row = table.add_row().cells
                row[0].text = pos["author"]
                row[1].text = pos["sentiment"]
                row[2].text = pos["text"][:120]

            doc.add_paragraph(
                "This requires a discussion before revisions can proceed.",
                style="Intense Quote",
            )

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_report.py <report.json> [--format markdown|docx] [--output file]")
        sys.exit(1)

    input_file = sys.argv[1]
    fmt = "markdown"
    output_file = None

    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            fmt = sys.argv[idx + 1]

    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_file = sys.argv[idx + 1]

    with open(input_file) as f:
        report = json.load(f)

    if fmt == "docx":
        out = output_file or "feedback_report.docx"
        render_docx(report, out)
        print(f"Report → {out}")
    else:
        md = render_markdown(report)
        if output_file:
            with open(output_file, "w") as f:
                f.write(md)
            print(f"Report → {output_file}")
        else:
            print(md)
